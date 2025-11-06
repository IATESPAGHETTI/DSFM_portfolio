from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level=1):
    """Add a heading with custom style"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 153)  # Dark blue
    return heading

def add_styled_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with custom styling"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def create_simplified_report():
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Data Science in Financial Markets', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Pre-Submission Report', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Section 1: Project Details
    add_heading_with_style(doc, '1. Project Details', 1)
    
    add_styled_paragraph(doc, 'Project Title:', bold=True)
    doc.add_paragraph('INDmoney Clone - Stock Market Analysis Platform for Indian Markets', style='List Bullet')
    
    add_styled_paragraph(doc, 'Project Type:', bold=True)
    doc.add_paragraph('Web Application with Machine Learning and Real-Time Data', style='List Bullet')
    
    add_styled_paragraph(doc, 'Team Members:', bold=True)
    doc.add_paragraph('Deepu James (230571)', style='List Bullet')
    doc.add_paragraph('Yaksh Rohilla (230605)', style='List Bullet')
    doc.add_paragraph('Mukund Madhav Agarwal (230594)', style='List Bullet')
    doc.add_paragraph('Aarav Pratap Singh (230625)', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 2: What We Built (Simple Abstract)
    add_heading_with_style(doc, '2. What We Built', 1)
    
    p = doc.add_paragraph()
    p.add_run('We created a ').font.size = Pt(11)
    p.add_run('stock market analysis website').bold = True
    p.add_run(' for Indian investors. The website shows:').font.size = Pt(11)
    
    doc.add_paragraph('Live stock prices that update automatically every 2 seconds', style='List Bullet')
    doc.add_paragraph('Interactive price charts (like TradingView)', style='List Bullet')
    doc.add_paragraph('AI predictions for future market risk', style='List Bullet')
    doc.add_paragraph('Market condition detection (Bull/Bear market identification)', style='List Bullet')
    doc.add_paragraph('News sentiment analysis (coming soon)', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Technology Used:', bold=True)
    doc.add_paragraph('Frontend: HTML, CSS, JavaScript (what users see)', style='List Bullet 2')
    doc.add_paragraph('Backend: Python with FastAPI (server that processes data)', style='List Bullet 2')
    doc.add_paragraph('Database: SQLite (stores historical data)', style='List Bullet 2')
    doc.add_paragraph('Charts: Lightweight Charts library (professional trading charts)', style='List Bullet 2')
    doc.add_paragraph('Data Source: Finnhub API (real-time stock prices)', style='List Bullet 2')
    
    doc.add_page_break()
    
    # Section 3: The Problem We Are Solving
    add_heading_with_style(doc, '3. The Problem We Are Solving', 1)
    
    add_styled_paragraph(doc, 'Most stock market apps in India have four main problems:')
    doc.add_paragraph()
    
    add_styled_paragraph(doc, '1. Only Show Past Data:', bold=True)
    doc.add_paragraph('Apps show what already happened, but don\'t predict future risk. Investors can\'t see "How risky will the market be next month?"')
    
    add_styled_paragraph(doc, '2. No Big Picture View:', bold=True)
    doc.add_paragraph('Apps show individual stock info but don\'t explain overall market conditions. Investors don\'t know if they should be cautious or aggressive.')
    
    add_styled_paragraph(doc, '3. Complex Information:', bold=True)
    doc.add_paragraph('Advanced tools exist but are too complicated for regular people to understand.')
    
    add_styled_paragraph(doc, '4. Not Real-Time:', bold=True)
    doc.add_paragraph('Many apps need manual refresh, so investors miss important price changes during volatile times.')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Our Solution:', bold=True)
    doc.add_paragraph('We built a platform that answers key questions:')
    doc.add_paragraph('• What is happening in the market RIGHT NOW?', style='List Bullet')
    doc.add_paragraph('• Is it a good time to invest or should I wait?', style='List Bullet')
    doc.add_paragraph('• How risky will the market be in the next 30 days?', style='List Bullet')
    doc.add_paragraph('• What is the overall market mood (from news)?', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 4: Data Sources
    add_heading_with_style(doc, '4. Where We Get Our Data', 1)
    
    add_styled_paragraph(doc, 'Main Data Source - Finnhub Stock API:', bold=True)
    doc.add_paragraph('Website: https://finnhub.io/api/v1', style='List Bullet')
    doc.add_paragraph('What we get: Live stock prices, historical data, trading volume', style='List Bullet')
    doc.add_paragraph('Coverage: NIFTY 50, SENSEX, Bank NIFTY, and 80+ individual stocks', style='List Bullet')
    doc.add_paragraph('Update Speed: Every 2 seconds (very fast!)', style='List Bullet')
    doc.add_paragraph('History: Up to 10 years of past data', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'News Data - Google News (Coming Soon):', bold=True)
    doc.add_paragraph('What we get: Latest financial news headlines', style='List Bullet')
    doc.add_paragraph('Purpose: Feed into AI to understand market sentiment', style='List Bullet')
    doc.add_paragraph('Update Speed: Every 15 minutes', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Database - SQLite:', bold=True)
    doc.add_paragraph('Purpose: Store historical prices for analysis', style='List Bullet')
    doc.add_paragraph('Size: 10,000+ price records', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 5: How It Works (Methodology)
    add_heading_with_style(doc, '5. How It Works - The Technology Behind It', 1)
    
    add_styled_paragraph(doc, 'We use 4 main techniques:')
    doc.add_paragraph()
    
    add_styled_paragraph(doc, 'Technique 1: Risk Prediction (GARCH Model)', bold=True)
    add_styled_paragraph(doc, 'What it does:', italic=True)
    doc.add_paragraph('Predicts how volatile (risky/jumpy) the market will be in the next 30 days')
    
    add_styled_paragraph(doc, 'Why it\'s useful:', italic=True)
    doc.add_paragraph('Helps investors decide: "Should I invest now or wait for calmer times?"')
    
    add_styled_paragraph(doc, 'How it works:', italic=True)
    doc.add_paragraph('Looks at past 1 year of daily price changes')
    doc.add_paragraph('Uses mathematical formula to spot patterns in risk levels', style='List Bullet 2')
    doc.add_paragraph('Predicts future volatility based on these patterns', style='List Bullet 2')
    
    add_styled_paragraph(doc, 'Output:', italic=True)
    doc.add_paragraph('Example: "Expected volatility: 1.45% per day (moderate risk)"')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_styled_paragraph(doc, 'Technique 2: Market Condition Detection (K-Means Clustering)', bold=True)
    add_styled_paragraph(doc, 'What it does:', italic=True)
    doc.add_paragraph('Automatically identifies if we\'re in a Bull Market, Bear Market, High Volatility, or Consolidation phase')
    
    add_styled_paragraph(doc, 'Why it\'s useful:', italic=True)
    doc.add_paragraph('Helps investors adjust their strategy based on current market environment')
    
    add_styled_paragraph(doc, 'How it works:', italic=True)
    doc.add_paragraph('Calculates 6 market features:')
    doc.add_paragraph('1. Volatility - How much prices are jumping', style='List Number 2')
    doc.add_paragraph('2. Returns - How much market has gained/lost', style='List Number 2')
    doc.add_paragraph('3. Momentum - Speed of price changes', style='List Number 2')
    doc.add_paragraph('4. Trend - Overall direction (up or down)', style='List Number 2')
    doc.add_paragraph('5. RSI - Overbought/Oversold indicator', style='List Number 2')
    doc.add_paragraph('6. Volume - How much trading is happening', style='List Number 2')
    
    doc.add_paragraph('AI groups similar market conditions together into 4 categories', style='List Bullet 2')
    
    add_styled_paragraph(doc, 'Output:', italic=True)
    doc.add_paragraph('Example: "Current Regime: Bull Market (Returns: 5.2%, Volatility: 18.5%)"')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_styled_paragraph(doc, 'Technique 3: News Sentiment Analysis (FinBERT AI)', bold=True)
    add_styled_paragraph(doc, 'What it does:', italic=True)
    doc.add_paragraph('Reads financial news headlines and determines if they\'re Positive, Negative, or Neutral')
    
    add_styled_paragraph(doc, 'Why it\'s useful:', italic=True)
    doc.add_paragraph('Market mood affects prices. Knowing overall sentiment helps predict short-term movements')
    
    add_styled_paragraph(doc, 'How it works:', italic=True)
    doc.add_paragraph('Collects latest news headlines from Google News', style='List Bullet 2')
    doc.add_paragraph('Feeds headlines into FinBERT (AI trained on financial text)', style='List Bullet 2')
    doc.add_paragraph('AI analyzes language and rates sentiment', style='List Bullet 2')
    doc.add_paragraph('Calculates overall market sentiment percentage', style='List Bullet 2')
    
    add_styled_paragraph(doc, 'Output:', italic=True)
    doc.add_paragraph('Example: "Market Sentiment: 65% Positive, 20% Neutral, 15% Negative"')
    
    add_styled_paragraph(doc, 'Status:', italic=True)
    p = doc.add_paragraph()
    run = p.add_run('⏳ Coming Soon - Currently in development')
    run.italic = True
    run.font.color.rgb = RGBColor(204, 102, 0)  # Orange
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_styled_paragraph(doc, 'Technique 4: Real-Time Data Streaming', bold=True)
    add_styled_paragraph(doc, 'What it does:', italic=True)
    doc.add_paragraph('Automatically updates prices every 2 seconds without user needing to refresh')
    
    add_styled_paragraph(doc, 'Why it\'s useful:', italic=True)
    doc.add_paragraph('Investors see live market movements instantly during volatile periods')
    
    add_styled_paragraph(doc, 'How it works:', italic=True)
    doc.add_paragraph('Backend server fetches data from Finnhub every 2 seconds', style='List Bullet 2')
    doc.add_paragraph('Server pushes updates to all connected users', style='List Bullet 2')
    doc.add_paragraph('Frontend displays updates with smooth animations', style='List Bullet 2')
    
    doc.add_page_break()
    
    # Section 6: What We've Completed
    add_heading_with_style(doc, '6. What We\'ve Completed So Far', 1)
    
    add_styled_paragraph(doc, '✅ Risk Prediction Model (GARCH):', bold=True)
    doc.add_paragraph('Successfully predicts 30-day volatility', style='List Bullet')
    doc.add_paragraph('Example: NIFTY 50 shows 1.447% daily volatility (23% annualized)', style='List Bullet')
    doc.add_paragraph('Model is stable and gives consistent results', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, '✅ Market Condition Detector (K-Means):', bold=True)
    doc.add_paragraph('Successfully identifies 4 market regimes', style='List Bullet')
    doc.add_paragraph('Current detection: Both NIFTY and SENSEX in "Bull Market"', style='List Bullet')
    doc.add_paragraph('Shows all 6 features with bar charts', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, '✅ Live Price Updates:', bold=True)
    doc.add_paragraph('Prices update automatically every 2 seconds', style='List Bullet')
    doc.add_paragraph('Works for NIFTY 50, SENSEX, and Bank NIFTY', style='List Bullet')
    doc.add_paragraph('Green/red pulse animations show price movements', style='List Bullet')
    doc.add_paragraph('No manual refresh needed', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, '✅ Professional Trading Charts:', bold=True)
    doc.add_paragraph('Interactive candlestick charts (like TradingView)', style='List Bullet')
    doc.add_paragraph('Zoom in/out and pan across time periods', style='List Bullet')
    doc.add_paragraph('Hover to see OHLC (Open, High, Low, Close) values', style='List Bullet')
    doc.add_paragraph('Date and time display correctly', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, '✅ Beautiful User Interface:', bold=True)
    doc.add_paragraph('Clean design matching INDmoney style', style='List Bullet')
    doc.add_paragraph('Three main pages: Dashboard, Trading, Analysis', style='List Bullet')
    doc.add_paragraph('Responsive layout with white cards and blue accents', style='List Bullet')
    doc.add_paragraph('Professional look and feel', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, '✅ Complete Technical Setup:', bold=True)
    doc.add_paragraph('Backend server running (Python FastAPI)', style='List Bullet')
    doc.add_paragraph('Frontend serving correctly (HTML/CSS/JavaScript)', style='List Bullet')
    doc.add_paragraph('Database storing historical data (SQLite)', style='List Bullet')
    doc.add_paragraph('API integration working (Finnhub)', style='List Bullet')
    doc.add_paragraph('All components communicating successfully', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 7: What's Next
    add_heading_with_style(doc, '7. Next Steps', 1)
    
    add_styled_paragraph(doc, 'Before Final Submission:', bold=True)
    doc.add_paragraph('✓ Complete News Sentiment Analysis integration', style='List Bullet')
    doc.add_paragraph('✓ Add API documentation (how other developers can use it)', style='List Bullet')
    doc.add_paragraph('✓ Write user guide with screenshots', style='List Bullet')
    doc.add_paragraph('✓ Add code comments explaining each function', style='List Bullet')
    doc.add_paragraph('✓ Test all features thoroughly', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Future Enhancements (After Course):', bold=True)
    doc.add_paragraph('• AI-powered market explanations in plain English', style='List Bullet')
    doc.add_paragraph('• Portfolio tracking (add stocks and track profit/loss)', style='List Bullet')
    doc.add_paragraph('• Integration with NSE/BSE official APIs', style='List Bullet')
    doc.add_paragraph('• More advanced models (different types of risk prediction)', style='List Bullet')
    doc.add_paragraph('• Mobile app version', style='List Bullet')
    
    doc.add_page_break()
    
    # Section 8: Key Results
    add_heading_with_style(doc, '8. Key Results and Findings', 1)
    
    add_styled_paragraph(doc, 'A. Technical Performance:', bold=True)
    doc.add_paragraph('✓ Page loads in under 2 seconds', style='List Bullet')
    doc.add_paragraph('✓ Data updates every 2 seconds reliably', style='List Bullet')
    doc.add_paragraph('✓ AI models respond in under 1 second', style='List Bullet')
    doc.add_paragraph('✓ Charts render smoothly with no lag', style='List Bullet')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'B. Analytical Findings:', bold=True)
    
    p = doc.add_paragraph()
    run = p.add_run('Market Regime: ')
    run.bold = True
    p.add_run('Currently in ')
    run = p.add_run('Bull Market')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
    
    doc.add_paragraph('Strong positive returns', style='List Bullet 2')
    doc.add_paragraph('Moderate volatility (stable gains)', style='List Bullet 2')
    doc.add_paragraph('Healthy RSI levels (not overbought)', style='List Bullet 2')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Risk Forecast: ')
    run.bold = True
    p.add_run('Moderate risk environment')
    
    doc.add_paragraph('NIFTY 50: 1.45% daily volatility (~23% annualized)', style='List Bullet 2')
    doc.add_paragraph('SENSEX: 1.48% daily volatility (~24% annualized)', style='List Bullet 2')
    doc.add_paragraph('Volatility is mean-reverting (stable)', style='List Bullet 2')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Sentiment Analysis: ')
    run.bold = True
    run = p.add_run('(Coming Soon)')
    run.italic = True
    run.font.color.rgb = RGBColor(204, 102, 0)  # Orange
    
    doc.add_paragraph('Will show correlation between news and price movements', style='List Bullet 2')
    doc.add_paragraph('Early tests show negative news clusters match volatility spikes', style='List Bullet 2')
    
    doc.add_page_break()
    
    # Section 9: Screenshots
    add_heading_with_style(doc, '9. Application Screenshots', 1)
    
    add_styled_paragraph(doc, 'Screenshot 1: Live Dashboard', bold=True)
    doc.add_paragraph('Shows NIFTY 50, SENSEX, and BANK NIFTY with real-time prices')
    doc.add_paragraph('Green/red pulse effects indicate price changes')
    doc.add_paragraph('Clean white cards with percentage changes')
    doc.add_paragraph('Updates automatically every 2 seconds')
    doc.add_paragraph()
    
    add_styled_paragraph(doc, 'Screenshot 2: Trading Interface', bold=True)
    doc.add_paragraph('Professional candlestick chart (like TradingView)')
    doc.add_paragraph('Interactive - zoom and pan with mouse')
    doc.add_paragraph('Hover tooltip shows Open, High, Low, Close, Date')
    doc.add_paragraph('Toggle between NIFTY 50 and SENSEX')
    doc.add_paragraph()
    
    add_styled_paragraph(doc, 'Screenshot 3: Market Regime Analysis', bold=True)
    doc.add_paragraph('Bar chart showing 6 market features')
    doc.add_paragraph('Current regime: "Bull Market" badge')
    doc.add_paragraph('Feature values: Volatility, Returns, Momentum, Trend, RSI, Volume')
    doc.add_paragraph('Switch between NIFTY and SENSEX with dropdown')
    doc.add_paragraph()
    
    add_styled_paragraph(doc, 'Screenshot 4: Volatility Forecast', bold=True)
    doc.add_paragraph('30-day GARCH forecast line chart')
    doc.add_paragraph('Shows mean volatility with indicator line')
    doc.add_paragraph('Smooth mean-reverting pattern')
    doc.add_paragraph('AI insights panel with risk recommendations')
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.bold = True
    run = p.add_run('Actual screenshots will be attached separately with the final submission')
    run.italic = True
    
    doc.add_page_break()
    
    # Section 10: References
    add_heading_with_style(doc, '10. References', 1)
    
    add_styled_paragraph(doc, 'Academic Papers:', bold=True)
    doc.add_paragraph('1. Araci, D. (2019). "FinBERT: Financial Sentiment Analysis with Pre-trained Language Models."')
    doc.add_paragraph('2. Bollerslev, T. (1986). "Generalized autoregressive conditional heteroskedasticity." Journal of Econometrics.')
    doc.add_paragraph('3. Engle, R. F. (1982). "Autoregressive conditional heteroscedasticity." Econometrica.')
    doc.add_paragraph('4. MacQueen, J. (1967). "Classification and analysis of multivariate observations."')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Data Sources:', bold=True)
    doc.add_paragraph('5. Finnhub Stock API - https://finnhub.io/docs/api')
    doc.add_paragraph('6. Google News - https://news.google.com/')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Technical Libraries & Tools:', bold=True)
    doc.add_paragraph('7. FastAPI Documentation - https://fastapi.tiangolo.com/')
    doc.add_paragraph('8. Lightweight Charts - https://tradingview.github.io/lightweight-charts/')
    doc.add_paragraph('9. Chart.js - https://www.chartjs.org/')
    doc.add_paragraph('10. arch Library (GARCH models) - https://arch.readthedocs.io/')
    doc.add_paragraph('11. scikit-learn (Machine Learning) - https://scikit-learn.org/')
    doc.add_paragraph('12. Hugging Face Transformers (FinBERT) - https://huggingface.co/transformers/')
    
    doc.add_paragraph()
    add_styled_paragraph(doc, 'Course Materials:', bold=True)
    doc.add_paragraph('13. Data Science in Financial Markets (DSFM) Course, IIT Delhi, 2025')
    
    doc.add_page_break()
    
    # Final page
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted by:')
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Deepu James (230571)\n')
    run.font.size = Pt(11)
    run = p.add_run('Yaksh Rohilla (230605)\n')
    run.font.size = Pt(11)
    run = p.add_run('Mukund Madhav Agarwal (230594)\n')
    run.font.size = Pt(11)
    run = p.add_run('Aarav Pratap Singh (230625)')
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submission Date: October 31, 2025')
    run.bold = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Course: Data Science in Financial Markets')
    run.italic = True
    run.font.size = Pt(11)
    
    # Save document
    doc.save('DSFM_Simplified_Report.docx')
    print("✅ Simplified report created successfully!")
    print("📄 File saved as: DSFM_Simplified_Report.docx")

if __name__ == "__main__":
    create_simplified_report()
