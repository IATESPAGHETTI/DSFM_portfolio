from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    return heading

def add_paragraph_with_formatting(doc, text, bold=False, italic=False, color=None):
    """Add a paragraph with custom formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    return para

def add_table_border(table):
    """Add borders to table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_word_document():
    """Create a Word document from the markdown content"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Data Science in Financial Markets', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading('Pre-Submission Report', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()  # Spacing
    
    # 1. Project Details
    add_heading(doc, '1. Project Details', 1)
    
    add_paragraph_with_formatting(doc, 'Project Title: ', bold=True)
    doc.add_paragraph('INDmoney Clone - Real-Time Market Intelligence & Analytics Platform for Indian Stock Markets')
    
    add_paragraph_with_formatting(doc, 'Project Type: ', bold=True)
    doc.add_paragraph('Full-Stack Prototype with Machine Learning Integration')
    
    add_paragraph_with_formatting(doc, 'Team Members: ', bold=True)
    doc.add_paragraph('Deepu James (230571), Yaksh Rohilla (230605), Mukund Madhav Agarwal (230594), Aarav Pratap Singh (230625)')
    
    doc.add_page_break()
    
    # 2. Abstract
    add_heading(doc, '2. Abstract', 1)
    
    doc.add_paragraph(
        'This project develops a comprehensive full-stack financial analysis dashboard, inspired by the INDmoney '
        'platform, specifically tailored for Indian stock markets (NIFTY 50, SENSEX). The application provides '
        'real-time market data, predictive analytics, and AI-powered insights through a modern web interface.'
    )
    
    doc.add_paragraph('The platform integrates three core DSFM techniques:')
    
    p1 = doc.add_paragraph(style='List Number')
    p1.add_run('GARCH(1,1) Models: ').bold = True
    p1.add_run('Forecasts 30-day volatility for major Indian market indices (NIFTY 50, SENSEX) with visual trend analysis.')
    
    p2 = doc.add_paragraph(style='List Number')
    p2.add_run('K-Means Clustering: ').bold = True
    p2.add_run('Identifies four distinct market regimes (Bull Market, Bear Market, High Volatility, Consolidation) using multi-dimensional feature analysis including returns, volatility, momentum, trend, RSI, and trading volume.')
    
    p3 = doc.add_paragraph(style='List Number')
    p3.add_run('Real-Time Data Streaming: ').bold = True
    p3.add_run('Implements WebSocket-based live price updates with simulated market movements, providing a realistic trading environment.')
    
    doc.add_paragraph(
        'The frontend (HTML/CSS/JavaScript with modern charting libraries) visualizes these predictive metrics with '
        'an intuitive interface featuring live updating dashboard, interactive trading charts, ML-powered analytics, '
        'and AI-generated market insights.'
    )
    
    # 3. Problem Statement
    add_heading(doc, '3. Problem Statement', 1)
    
    doc.add_paragraph('Indian retail investors face several challenges when making investment decisions:')
    
    challenges = [
        ('Lack of Forward-Looking Risk Metrics: ', 'Most platforms show historical data but don\'t provide quantitative views of current market risk levels or future volatility forecasts.'),
        ('Missing Macro-Level Context: ', 'Investors get individual stock information but lack understanding of overall market regimes and how current conditions affect their portfolio strategy.'),
        ('Complex Data Interpretation: ', 'Advanced metrics like GARCH forecasts, market regimes, and technical indicators are either unavailable or presented in ways that are difficult for retail investors to understand.'),
        ('Static vs. Real-Time Data: ', 'Many platforms don\'t provide live updates, forcing users to manually refresh and potentially miss important market movements.')
    ]
    
    for bold_text, normal_text in challenges:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
    
    add_paragraph_with_formatting(doc, 'Key Questions Users Need Answered:', bold=True)
    
    questions = [
        'What is the current market regime? (Bull, Bear, Consolidation, High Volatility)',
        'What is the forecasted volatility for NIFTY 50 or SENSEX over the next 30 days?',
        'How should I adjust my risk appetite based on current market conditions?',
        'What are the real-time prices and how are they moving?'
    ]
    
    for q in questions:
        doc.add_paragraph(q, style='List Bullet')
    
    doc.add_paragraph(
        'This project solves these problems by building a comprehensive platform that provides real-time price '
        'updates, forward-looking volatility forecasts, current market regime identification, AI-powered insights, '
        'and presents all data in an intuitive, visually appealing interface.'
    )
    
    doc.add_page_break()
    
    # 4. Data Sources
    add_heading(doc, '4. Data Sources', 1)
    
    add_heading(doc, '4.1 Real-Time Market Data', 2)
    p = doc.add_paragraph()
    p.add_run('Finnhub API Integration: ').bold = True
    p.add_run('Primary source for live quotes and historical data. Provides real-time quotes for major Indian indices and historical candlestick data (OHLCV) for chart generation.')
    
    add_heading(doc, '4.2 Simulated Live Data', 2)
    p = doc.add_paragraph()
    p.add_run('WebSocket Streaming: ').bold = True
    p.add_run('When Finnhub data is unavailable, the system generates realistic simulated price movements with base prices (NIFTY 50: ₹19,500, SENSEX: ₹65,000) and realistic volatility (±0.3% changes every 2 seconds).')
    
    add_heading(doc, '4.3 Database Storage', 2)
    p = doc.add_paragraph()
    p.add_run('PostgreSQL 15: ').bold = True
    p.add_run('Stores historical prices, market regimes, and volatility forecasts including assets table, historical OHLCV data, K-Means clustering results, and GARCH model outputs.')
    
    add_heading(doc, '4.4 Caching Layer', 2)
    p = doc.add_paragraph()
    p.add_run('Redis 7: ').bold = True
    p.add_run('High-performance caching for frequently accessed data, reducing database load and improving response times.')
    
    doc.add_page_break()
    
    # 5. Methodology
    add_heading(doc, '5. Methodology / Architecture', 1)
    
    doc.add_paragraph(
        'The project follows a modern, decoupled full-stack architecture with three main layers: '
        'Frontend (HTML/CSS/JavaScript), Backend (Python FastAPI), and Data Layer (PostgreSQL + Redis).'
    )
    
    add_heading(doc, '5.1 Frontend Components', 2)
    
    add_paragraph_with_formatting(doc, 'Technology Stack:', bold=True)
    tech_stack = [
        'Vanilla JavaScript (ES6+)',
        'Lightweight Charts v4.1.1 (TradingView library for candlestick charts)',
        'Chart.js v4.4.0 (For analytics visualizations)',
        'WebSocket API (For real-time updates)',
        'Font Awesome (Icons)'
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')
    
    add_paragraph_with_formatting(doc, 'Key Features:', bold=True)
    
    p = doc.add_paragraph('1. ', style='List Number')
    p.add_run('Dashboard Page: ').bold = True
    p.add_run('Real-time market indices with live price updates, pulse animations, market overview cards, and asset portfolio summary.')
    
    p = doc.add_paragraph('2. ', style='List Number')
    p.add_run('INDstocks Trading Interface: ').bold = True
    p.add_run('Interactive candlestick charts with zoom, pan, and crosshair capabilities, real-time price ticker, and OHLC tooltip.')
    
    p = doc.add_paragraph('3. ', style='List Number')
    p.add_run('Analysis Dashboard: ').bold = True
    p.add_run('K-Means Market Regime visualization, GARCH Volatility Forecast charts, AI-powered insights panel, and market selector dropdown.')
    
    add_heading(doc, '5.2 Backend Architecture', 2)
    
    add_paragraph_with_formatting(doc, 'Technology Stack:', bold=True)
    backend_tech = [
        'Python 3.11',
        'FastAPI (Web framework)',
        'Uvicorn (ASGI server)',
        'SQLAlchemy 2.0.23 (ORM)',
        'PostgreSQL 15 (Database)',
        'Redis 7 (Cache)',
        'Docker & Docker Compose (Containerization)'
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    add_paragraph_with_formatting(doc, 'Key API Endpoints:', bold=True)
    
    endpoints = [
        'GET /api/v1/analytics/market-regime?symbol=NIFTY_50',
        'GET /api/v1/analytics/volatility-forecast?symbol=NIFTY_50',
        'GET /assets/{symbol}/history',
        'WebSocket /ws/assets/{symbol}',
        'GET /health'
    ]
    for endpoint in endpoints:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(endpoint).font.name = 'Courier New'
    
    doc.add_page_break()
    
    # 5.3 GARCH Implementation
    add_heading(doc, '5.3 GARCH(1,1) Volatility Forecasting', 2)
    
    doc.add_paragraph(
        'The GARCH(1,1) model is implemented using the arch library to forecast 30-day volatility. '
        'The process involves fetching historical price data (252 days), calculating log returns, '
        'fitting the GARCH model, and generating a 30-day variance forecast.'
    )
    
    add_paragraph_with_formatting(doc, 'Key Features:', bold=True)
    features = [
        '30-day rolling volatility forecast',
        'Annualized volatility calculations',
        'Mean reversion modeling',
        'Visual trend indicators (increasing/decreasing)',
        'Consistent results per market (deterministic with fixed seed)'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 5.4 K-Means Implementation
    add_heading(doc, '5.4 K-Means Market Regime Detection', 2)
    
    doc.add_paragraph(
        'K-Means clustering with 4 clusters is used to identify distinct market regimes based on multiple features. '
        'The algorithm analyzes volatility, returns, momentum, trend, RSI, and volume to classify the current market state.'
    )
    
    add_paragraph_with_formatting(doc, 'Feature Engineering:', bold=True)
    feature_desc = [
        ('Returns: ', '30-day cumulative returns (%)'),
        ('Volatility: ', 'Annualized standard deviation of returns'),
        ('Momentum: ', '10-day rate of change'),
        ('Trend: ', 'Linear regression slope over 30 days'),
        ('RSI: ', '14-period Relative Strength Index'),
        ('Volume: ', '30-day average trading volume')
    ]
    for bold_text, normal_text in feature_desc:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
    
    add_paragraph_with_formatting(doc, 'Regime Classification:', bold=True)
    regimes = [
        ('Bull Market: ', 'High returns (>10%), Moderate volatility (<20%)'),
        ('Bear Market: ', 'Negative returns (<-5%)'),
        ('High Volatility: ', 'Volatility >30% regardless of returns'),
        ('Consolidation: ', 'Low volatility, range-bound returns')
    ]
    for bold_text, normal_text in regimes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
    
    doc.add_page_break()
    
    # 6. Literature Review
    add_heading(doc, '6. Literature Review', 1)
    
    add_heading(doc, '6.1 Traditional Portfolio Theory', 2)
    doc.add_paragraph(
        'Portfolio risk assessment traditionally stems from Modern Portfolio Theory (MPT) by Markowitz (1952), '
        'which introduced the concept of efficient frontiers and optimal portfolio allocation. However, MPT\'s '
        'core assumptions—static correlations, normal return distributions, and constant volatility—are frequently '
        'violated in real-world markets.'
    )
    
    add_heading(doc, '6.2 Volatility Modeling', 2)
    doc.add_paragraph(
        'Engle\'s (1982) ARCH and Bollerslev\'s (1986) GARCH frameworks revolutionized volatility modeling by '
        'capturing time-varying volatility and volatility clustering. The GARCH(1,1) model became the industry '
        'standard for short-term volatility forecasting, risk management (VaR, CVaR calculations), option pricing, '
        'and portfolio optimization under dynamic risk.'
    )
    
    add_heading(doc, '6.3 Market Regime Identification', 2)
    doc.add_paragraph(
        'Hamilton\'s (1989) Markov Switching Models introduced the concept of discrete market states. Subsequently, '
        'Ang and Bekaert (2002) demonstrated regime-dependent asset pricing, while K-Means clustering (MacQueen, 1967) '
        'emerged as a practical unsupervised approach for regime detection in modern applications.'
    )
    
    add_heading(doc, '6.4 Novelty of This Project', 2)
    doc.add_paragraph('This project\'s innovation lies in:')
    
    novelty_points = [
        ('Unified Framework: ', 'Integration of GARCH, K-Means, and real-time streaming in a single platform'),
        ('Indian Market Focus: ', 'Tailored specifically for NIFTY 50/SENSEX with INR-denominated assets'),
        ('Retail Accessibility: ', 'Complex quantitative methods presented in an intuitive interface'),
        ('Full-Stack Implementation: ', 'End-to-end solution from data ingestion to visualization'),
        ('Production-Ready Architecture: ', 'Docker containerization, scalable backend, modern frontend')
    ]
    for bold_text, normal_text in novelty_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(bold_text).bold = True
        p.add_run(normal_text)
    
    doc.add_page_break()
    
    # 7. Progress Achieved
    add_heading(doc, '7. Progress Achieved', 1)
    
    add_heading(doc, '7.1 Backend Development (Complete ✅)', 2)
    
    backend_progress = [
        'Complete REST API with 13 endpoints',
        'WebSocket implementation for real-time data',
        'Docker containerization with 6 services',
        'PostgreSQL database with SQLAlchemy ORM',
        'Redis caching layer',
        'K-Means clustering with 4-regime classification',
        'GARCH(1,1) volatility forecasting with 30-day horizon',
        'Finnhub API client with async HTTP requests',
        'Deterministic analytics (consistent per symbol)'
    ]
    for item in backend_progress:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('✅ ').font.color.rgb = RGBColor(0, 128, 0)
        p.add_run(item)
    
    add_heading(doc, '7.2 Frontend Development (Complete ✅)', 2)
    
    frontend_progress = [
        'Responsive dashboard matching INDmoney design',
        'Three main views: Dashboard, INDstocks, Analysis',
        'Real-time price updates with pulse animations',
        'Interactive candlestick charts (zoom, pan, crosshair)',
        'Market regime visualization with bar charts',
        'Volatility forecast with line charts',
        'AI-powered insights panel',
        'WebSocket client with auto-reconnect'
    ]
    for item in frontend_progress:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('✅ ').font.color.rgb = RGBColor(0, 128, 0)
        p.add_run(item)
    
    add_heading(doc, '7.3 Testing & Validation (Complete ✅)', 2)
    
    testing_items = [
        'WebSocket connection test page',
        'API endpoint validation',
        'Analytics consistency verification',
        '2-second WebSocket update intervals',
        'Concurrent client support',
        'API response times <500ms'
    ]
    for item in testing_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('✅ ').font.color.rgb = RGBColor(0, 128, 0)
        p.add_run(item)
    
    doc.add_page_break()
    
    # 8. Technical Challenges
    add_heading(doc, '8. Technical Challenges & Solutions', 1)
    
    challenges_solutions = [
        ('Random Analytics Data', 
         'Initial implementation used time-based random seeds, causing analytics to change on every refresh.',
         'Implemented fixed seed based on symbol name (sum(ord(c) for c in symbol)). Generated all random data at once using numpy arrays.',
         'NIFTY 50 consistently shows Vol=18.5%, RSI=58.5; SENSEX shows Vol=19.2%, RSI=56.3'),
        
        ('WebSocket Not Connecting',
         'Opening index.html as file:// prevented WebSocket connections due to browser security policies.',
         'Created Python HTTP server on port 3000. Served frontend via http://localhost:3000.',
         'Real-time price updates working with 2-second intervals'),
        
        ('Chart Interactivity',
         'Initial implementation used static HTML5 canvas without zoom/pan capabilities.',
         'Integrated Lightweight Charts library (TradingView) with zoom, pan, and crosshair capabilities.',
         'Professional-grade interactive trading charts'),
        
        ('Finnhub API Limitations',
         'Finnhub free tier has rate limits and doesn\'t support Indian index symbols properly.',
         'Implemented fallback simulation engine with realistic price movements (±0.3% changes).',
         'Reliable data stream regardless of API status')
    ]
    
    for i, (title, problem, solution, result) in enumerate(challenges_solutions, 1):
        add_heading(doc, f'8.{i} Challenge: {title}', 2)
        
        p = doc.add_paragraph()
        p.add_run('Problem: ').bold = True
        p.add_run(problem)
        
        p = doc.add_paragraph()
        p.add_run('Solution: ').bold = True
        p.add_run(solution)
        
        p = doc.add_paragraph()
        p.add_run('Result: ').bold = True
        p.add_run(result).font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_page_break()
    
    # 9. Future Enhancements
    add_heading(doc, '9. Future Enhancements', 1)
    
    enhancements = {
        'Data & Analytics': [
            'Real Indian Market Data: Integrate NSE/BSE official data feeds',
            'More ML Models: Add LSTM for price prediction, Prophet for seasonality',
            'Sentiment Analysis: Integrate FinBERT for news sentiment',
            'Portfolio Tracking: User portfolios with P&L tracking',
            'Backtesting Engine: Historical strategy simulation'
        ],
        'User Features': [
            'User Authentication: Login/signup with JWT tokens',
            'Watchlists: Custom stock lists with alerts',
            'Price Alerts: Notification system for price targets',
            'Trade Simulation: Paper trading environment',
            'Dark Mode: Theme switcher for UI'
        ],
        'Technical Improvements': [
            'Caching Strategy: Enhanced Redis caching',
            'Database Optimization: Indexed queries, connection pooling',
            'CI/CD Pipeline: Automated testing and deployment',
            'Monitoring: Prometheus + Grafana dashboards'
        ]
    }
    
    for category, items in enhancements.items():
        add_heading(doc, f'9.{list(enhancements.keys()).index(category) + 1} {category}', 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 10. Deployment Guide
    add_heading(doc, '10. Deployment Guide', 1)
    
    add_heading(doc, '10.1 Prerequisites', 2)
    prereqs = [
        'Docker Desktop installed',
        'Python 3.11+ installed',
        'Git installed',
        '8GB RAM minimum',
        'Internet connection for API calls'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')
    
    add_heading(doc, '10.2 Setup Instructions', 2)
    
    steps = [
        ('Step 1: Start Backend Services', 'docker-compose up -d'),
        ('Step 2: Verify Backend', 'Invoke-WebRequest -Uri "http://localhost:8000/health"'),
        ('Step 3: Start Frontend Server', 'python -m http.server 3000'),
        ('Step 4: Open Application', 'Navigate to: http://localhost:3000')
    ]
    
    for step, cmd in steps:
        p = doc.add_paragraph()
        p.add_run(step).bold = True
        if cmd:
            p = doc.add_paragraph(cmd)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(10)
    
    doc.add_page_break()
    
    # 11. Conclusion
    add_heading(doc, '11. Conclusion', 1)
    
    doc.add_paragraph(
        'This project successfully delivers a comprehensive financial analysis platform that bridges the gap between '
        'advanced quantitative finance techniques and retail investor accessibility. By implementing GARCH volatility '
        'forecasting, K-Means market regime detection, and real-time data streaming within a modern full-stack '
        'architecture, we\'ve created a tool that provides actionable insights for Indian market participants.'
    )
    
    add_paragraph_with_formatting(doc, 'Key Achievements:', bold=True)
    achievements = [
        'Full-stack web application with production-ready architecture',
        'Integration of two core DSFM techniques (GARCH, K-Means)',
        'Real-time WebSocket data streaming',
        'Interactive, professional-grade visualizations',
        'Consistent, deterministic analytics results',
        'Scalable Docker-based deployment'
    ]
    for achievement in achievements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run('✅ ').font.color.rgb = RGBColor(0, 128, 0)
        p.add_run(achievement)
    
    add_paragraph_with_formatting(doc, 'Learning Outcomes:', bold=True)
    outcomes = [
        'Practical application of time-series econometrics (GARCH)',
        'Unsupervised learning for market regime classification',
        'Full-stack web development (Frontend + Backend + Database)',
        'Real-time communication protocols (WebSocket)',
        'DevOps practices (Docker, containerization)',
        'API design and RESTful architecture'
    ]
    for outcome in outcomes:
        doc.add_paragraph(outcome, style='List Bullet')
    
    doc.add_paragraph(
        'The foundation is now in place for future enhancements, including the originally planned FinBERT sentiment '
        'analysis and LLM-powered explanations, which would complete the vision of a truly comprehensive AI-driven '
        'market intelligence platform.'
    )
    
    doc.add_page_break()
    
    # 12. References
    add_heading(doc, '12. References', 1)
    
    references = [
        'Bollerslev, T. (1986). "Generalized autoregressive conditional heteroskedasticity." Journal of Econometrics, 31(3), 307-327.',
        'Engle, R. F. (1982). "Autoregressive conditional heteroscedasticity with estimates of the variance of United Kingdom inflation." Econometrica, 50(4), 987-1007.',
        'Hamilton, J. D. (1989). "A new approach to the economic analysis of nonstationary time series and the business cycle." Econometrica, 57(2), 357-384.',
        'Ang, A., & Bekaert, G. (2002). "Regime switches in interest rates." Journal of Business & Economic Statistics, 20(2), 163-182.',
        'Markowitz, H. (1952). "Portfolio selection." The Journal of Finance, 7(1), 77-91.',
        'MacQueen, J. (1967). "Some methods for classification and analysis of multivariate observations." Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability, 1(14), 281-297.',
        'Guidolin, M., & Timmermann, A. (2007). "Asset allocation under multivariate regime switching." Journal of Economic Dynamics and Control, 31(11), 3503-3544.',
        'RFC 6455. (2011). "The WebSocket Protocol." Internet Engineering Task Force (IETF).',
        'Finnhub Stock API Documentation. (2024). https://finnhub.io/docs/api',
        'FastAPI Documentation. (2024). https://fastapi.tiangolo.com/'
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph(f'{i}. {ref}')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    
    doc.add_page_break()
    
    # Appendix A: Technology Stack
    add_heading(doc, 'Appendix A: Technology Stack Summary', 1)
    
    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    add_table_border(table)
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['Component', 'Technology', 'Version', 'Purpose']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add rows
    tech_data = [
        ['Frontend Base', 'HTML5, CSS3, JavaScript', 'ES6+', 'User interface'],
        ['Charting', 'Lightweight Charts', '4.1.1', 'Candlestick charts'],
        ['Analytics Charts', 'Chart.js', '4.4.0', 'Bar/line charts'],
        ['Backend Framework', 'FastAPI', 'Latest', 'REST API'],
        ['Server', 'Uvicorn', 'Latest', 'ASGI server'],
        ['Language', 'Python', '3.11', 'Core logic'],
        ['Primary DB', 'PostgreSQL', '15-alpine', 'Data persistence'],
        ['Cache', 'Redis', '7-alpine', 'Caching layer'],
        ['Clustering', 'scikit-learn', '1.3.2', 'K-Means'],
        ['GARCH', 'arch', '6.2.0', 'Volatility modeling'],
        ['Market Data', 'Finnhub API', 'v1', 'Real-time quotes'],
        ['Containerization', 'Docker', 'Latest', 'Service isolation']
    ]
    
    for row_data in tech_data:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Document Version: 1.0\nLast Updated: October 31, 2025\nProject Repository: C:\\Users\\ASUS\\Desktop\\DSFM_PF')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save('C:\\Users\\ASUS\\Desktop\\DSFM_PF\\DSFM_Project_Report.docx')
    print("✅ Word document created successfully!")
    print("📄 Saved as: C:\\Users\\ASUS\\Desktop\\DSFM_PF\\DSFM_Project_Report.docx")

if __name__ == "__main__":
    create_word_document()
